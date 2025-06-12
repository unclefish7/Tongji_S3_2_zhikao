import json
import sys
import os
import shutil
import tempfile
import secrets
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup, NavigableString
from bs4.element import Tag
from matplotlib import rcParams
from collections import defaultdict
import base64
import subprocess
from PIL import Image
import fitz  # PyMuPDF

# 加密配置 - 与Node.js保持一致
ENCRYPTION_KEY = b'zhikao-system-2024-secure-key-32b'

def get_key():
    """获取标准化的密钥"""
    if len(ENCRYPTION_KEY) < 32:
        return ENCRYPTION_KEY.ljust(32, b'\x00')
    elif len(ENCRYPTION_KEY) > 32:
        return ENCRYPTION_KEY[:32]
    return ENCRYPTION_KEY

def decrypt_data(encrypted_data):
    """解密数据 - 与Node.js保持一致"""
    try:
        parts = encrypted_data.split(':')
        if len(parts) != 2:
            return None
            
        iv = bytes.fromhex(parts[0])
        encrypted_bytes = bytes.fromhex(parts[1])
        
        key_bytes = get_key()[:16]
        decrypted = bytearray()
        
        for i, byte in enumerate(encrypted_bytes):
            decrypted.append(byte ^ key_bytes[i % len(key_bytes)] ^ iv[i % len(iv)])
        
        return decrypted.decode('utf-8')
    except Exception as e:
        print(f"解密失败: {e}")
        return None

def is_encrypted_format(content):
    """检查内容是否是加密格式"""
    return (':' in content and 
            not content.strip().startswith('{') and 
            not content.strip().startswith('['))

def read_encrypted_file(file_path):
    """读取并解密文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read().strip()
        
        if is_encrypted_format(content):
            decrypted_content = decrypt_data(content)
            if decrypted_content:
                return json.loads(decrypted_content)
            else:
                raise ValueError("解密失败")
        else:
            # 向后兼容：如果不是加密格式，直接解析JSON
            return json.loads(content)
    except Exception as e:
        print(f"读取加密文件失败: {e}")
        raise

# 设置中文字体支持
rcParams['font.sans-serif'] = ['SimHei']
rcParams['axes.unicode_minus'] = False

DEFAULT_FONT_NAME = '宋体'
ENGLISH_FONT_NAME = 'Times New Roman'
DEFAULT_FORMULA_HEIGHT_PT = 14

def generate_latex_image(latex_content, image_path):
    try:
        if not latex_content.strip():
            return None
        latex_content = latex_content.replace('\\pix', r'\pi x')
        fig, ax = plt.subplots()
        ax.axis('off')
        text = ax.text(0.5, 0.5, f'${latex_content}$', fontsize=DEFAULT_FORMULA_HEIGHT_PT,
                       ha='center', va='center')
        fig.canvas.draw()
        renderer = fig.canvas.get_renderer()
        bbox = text.get_window_extent(renderer=renderer)
        width_inch = bbox.width / fig.dpi
        height_inch = bbox.height / fig.dpi
        fig.set_size_inches(max(width_inch, 0.5), max(height_inch, 0.3))

        os.makedirs(image_path, exist_ok=True)
        image_file_path = os.path.join(image_path, f'latex_{abs(hash(latex_content))}.png')
        plt.savefig(image_file_path, dpi=300, bbox_inches='tight', pad_inches=0.05)
        plt.close(fig)
        return image_file_path
    except Exception as e:
        print(f"❌ 生成公式图片失败: {e}")
        return None

def set_font(run, is_title=False):
    run.font.name = DEFAULT_FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    run._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT_NAME)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT_NAME)
    run.font.size = Pt(16 if is_title else 12)
    run.bold = is_title

def insert_table(doc, html):
    soup = BeautifulSoup(html, 'html.parser')
    table_tag = soup.find('table')
    if not table_tag:
        return
    rows = table_tag.find_all('tr')
    table = doc.add_table(rows=0, cols=len(rows[0].find_all(['td', 'th'])))
    
    # 设置表格样式为带边框的表格
    table.style = 'Table Grid'
    
    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_cells = table.add_row().cells
        for i, cell in enumerate(cells):
            row_cells[i].text = cell.get_text(strip=True)
            para = row_cells[i].paragraphs[0]
            if para.runs:
                set_font(para.runs[0])

def process_rich_text(doc, content, image_dir, question_number=None):
    soup = BeautifulSoup(content, 'html.parser')
    first_paragraph = True
    
    for block in soup.find_all(['p', 'table', 'br']):
        if block.name == 'p':
            p = doc.add_paragraph()
            
            # 如果是第一个段落且有题目序号，先添加序号
            if first_paragraph and question_number is not None:
                number_run = p.add_run(f"{question_number}、")
                set_font(number_run)
                number_run.bold = True
                first_paragraph = False
            
            for child in block.children:
                if isinstance(child, NavigableString):
                    if child.strip():
                        run = p.add_run(str(child))
                        set_font(run)
                elif isinstance(child, Tag):
                    if child.name == 'span' and child.get('data-w-e-type') == 'formula':
                        latex_code = child.get('data-value', '')
                        img_path = generate_latex_image(latex_code, image_dir)
                        if img_path:
                            p.add_run().add_picture(img_path, height=Inches(0.2))
                    elif child.name == 'img':
                        # 处理编辑器中的图片
                        img_src = child.get('src', '')
                        if img_src:
                            processed_img_path = process_image_src(img_src)
                            if processed_img_path and os.path.exists(processed_img_path):
                                try:
                                    # 获取图片尺寸信息
                                    width_style = child.get('style', '')
                                    img_width = extract_image_width(width_style)
                                    p.add_run().add_picture(processed_img_path, width=img_width)
                                except Exception as e:
                                    print(f"插入图片失败: {e}")
                                    # 插入错误提示文本
                                    run = p.add_run("[图片加载失败]")
                                    set_font(run)
                    elif child.name == 'span':
                        run = p.add_run(child.get_text())
                        set_font(run)
        elif block.name == 'table':
            doc.add_paragraph()
            insert_table(doc, str(block))
            doc.add_paragraph()
        elif block.name == 'br':
            doc.add_paragraph()

def get_exam_metadata(input_file_path):
    """
    根据输入文件路径获取试卷元信息
    """
    try:
        # 获取输入文件名（不含扩展名）作为试卷ID
        paper_id = os.path.splitext(os.path.basename(input_file_path))[0]
        
        # 构建 totalExam.json 的路径
        # 目录结构: 项目根/zhikao/python/matplotlibphoto.py 和 项目根/data/exam/totalExam.json
        if getattr(sys, 'frozen', False):  # PyInstaller 打包后
            exe_dir = os.path.dirname(sys.executable)
            if os.path.basename(exe_dir) == 'python':
                # 从 项目根/zhikao/python 回到 项目根
                project_root = os.path.dirname(os.path.dirname(exe_dir))
            else:
                # 如果exe在项目根目录
                project_root = exe_dir
        else:
            # 脚本路径: 项目根/zhikao/python/matplotlibphoto.py
            # 需要回到项目根: 项目根
            script_dir = os.path.dirname(os.path.abspath(__file__))  # 项目根/zhikao/python
            zhikao_dir = os.path.dirname(script_dir)  # 项目根/zhikao
            project_root = os.path.dirname(zhikao_dir)  # 项目根
        
        total_exam_path = os.path.join(project_root, "data", "exam", "totalExam.json")
        print(f"查找试卷元信息文件: {total_exam_path}")
        
        if not os.path.exists(total_exam_path):
            print(f"⚠️ 试卷元信息文件不存在: {total_exam_path}")
            return None
        
        # 读取并解密 totalExam.json
        exam_list = read_encrypted_file(total_exam_path)
        
        # 查找匹配的试卷ID
        for exam in exam_list:
            if exam.get('paperId') == paper_id:
                print(f"✅ 找到匹配的试卷元信息: {paper_id}")
                return exam
        
        print(f"⚠️ 未找到匹配的试卷ID: {paper_id}")
        return None
        
    except Exception as e:
        print(f"❌ 获取试卷元信息失败: {e}")
        return None

def generate_docx(questions, output_path, input_file_path=None, mode='normal'):
    # 添加调试信息：打印试卷内容
    print("=" * 50)
    if mode == 'answer':
        print("开始生成答案，试卷内容如下：")
    else:
        print("开始生成试卷，试卷内容如下：")
    print("=" * 50)
    
    for i, question in enumerate(questions, 1):
        print(f"\n题目 {i}:")
        print(f"  ID: {question.get('id', '未知')}")
        print(f"  类型: {question.get('type', '未知')}")
        print(f"  分数: {question.get('score', '未知')}")
        print(f"  富文本内容: {question.get('richTextContent', '无内容')[:200]}...")
        
        # 检查是否包含图片
        content = question.get('richTextContent', '')
        if '<img' in content:
            print("  ⚠️ 发现图片标签")
            # 提取图片src
            soup = BeautifulSoup(content, 'html.parser')
            images = soup.find_all('img')
            for img in images:
                src = img.get('src', '')
                print(f"    图片源: {src}")
        
        # 检查是否包含公式
        if 'data-w-e-type="formula"' in content:
            print("  📐 发现数学公式")
            soup = BeautifulSoup(content, 'html.parser')
            formulas = soup.find_all('span', {'data-w-e-type': 'formula'})
            for formula in formulas:
                latex_code = formula.get('data-value', '')
                print(f"    公式内容: {latex_code}")
    
    print("=" * 50)
    print("开始处理文档生成...")
    print("=" * 50)
    
    doc = Document()
    
    # 添加试卷元信息（如果提供了输入文件路径）
    if input_file_path:
        exam_metadata = get_exam_metadata(input_file_path)
        if exam_metadata:
            # 试卷标题
            title_para = doc.add_paragraph()
            title_text = exam_metadata.get('name', '试卷')
            if mode == 'answer':
                title_text += " - 答案"
            title_run = title_para.add_run(title_text)
            set_font(title_run, is_title=True)
            title_run.font.size = Pt(20)
            title_para.alignment = 1  # 居中对齐
            
            # 试卷信息
            info_para = doc.add_paragraph()
            info_text = f"总分：{exam_metadata.get('score', '未知')}分    " \
                       f"部门：{exam_metadata.get('department', '未知')}    " \
                       f"考试时长：{exam_metadata.get('duration', '未知')}"
            info_run = info_para.add_run(info_text)
            set_font(info_run)
            info_para.alignment = 1  # 居中对齐
            
            # 添加分隔线
            doc.add_paragraph("=" * 60).alignment = 1
            doc.add_paragraph()
    
    type_order = ['选择题', '判断题', '填空题', '主观题']
    type_map = defaultdict(list)
    for q in questions:
        type_map[q['type']].append(q)

    image_dir = tempfile.mkdtemp(prefix='latex_images_')
    print(f"临时图片目录: {image_dir}")

    try:
        # 题目总序号计数器
        question_number = 1
        # 题型序号映射
        type_chinese_numbers = ['一', '二', '三', '四']
        type_index = 0
        
        for qtype in type_order:
            qlist = type_map[qtype]
            if not qlist:
                continue
            total = sum(q.get('score', 0) for q in qlist)
            avg = total / len(qlist) if qlist else 0
            
            # 添加题型标题，包含中文序号
            title = f"{type_chinese_numbers[type_index]}、{qtype}（共{len(qlist)}题，合计{total}分，每题{avg:.1f}分）"
            run = doc.add_paragraph().add_run(title)
            set_font(run, is_title=True)
            print(f"处理题目类型: {qtype}, 共 {len(qlist)} 题")
            type_index += 1

            for q in qlist:
                print(f"  处理题目 ID {q.get('id', '未知')}, 序号: {question_number}")
                
                # 添加题目前的换行
                doc.add_paragraph()
                
                # 处理题目内容，传入题目序号
                process_rich_text(doc, q.get('richTextContent', ''), image_dir, question_number)
                
                # 如果是答案模式，添加答案部分
                if mode == 'answer':
                    answer_content = q.get('answer', '')
                    if answer_content and answer_content.strip() and answer_content != '<p></p>':
                        # 添加答案标题
                        answer_title_para = doc.add_paragraph()
                        answer_title_run = answer_title_para.add_run("答案：")
                        set_font(answer_title_run)
                        answer_title_run.bold = True
                        answer_title_run.font.color.rgb = None  # 使用默认颜色
                        
                        # 添加答案内容
                        process_rich_text(doc, answer_content, image_dir)
                        
                        # 添加答案分隔线
                        separator_para = doc.add_paragraph()
                        separator_run = separator_para.add_run("─" * 50)
                        set_font(separator_run)
                        separator_run.font.size = Pt(8)
                    else:
                        # 如果没有答案，添加提示
                        no_answer_para = doc.add_paragraph()
                        no_answer_run = no_answer_para.add_run("答案：（未提供答案）")
                        set_font(no_answer_run)
                        no_answer_run.italic = True
                
                # 为主观题添加额外空行
                if qtype == '主观题' and mode != 'answer':
                    for _ in range(3):
                        doc.add_paragraph()
                
                question_number += 1

        doc.save(output_path)
        if mode == 'answer':
            print(f"✅ Word 答案已生成：{output_path}")
        else:
            print(f"✅ Word 试卷已生成：{output_path}")
    finally:
        print(f"清理临时目录: {image_dir}")
        shutil.rmtree(image_dir, ignore_errors=True)

def process_image_src(img_src):
    """
    处理图片源路径，支持多种格式
    """
    print(f"🖼️ 开始处理图片路径: {img_src}")
    
    try:
        # 处理 file:/// 协议的路径
        if img_src.startswith('file:///'):
            # 移除 file:/// 前缀
            file_path = img_src[8:]  # 去掉 'file:///'
            # 获取项目根目录
            if getattr(sys, 'frozen', False):  # PyInstaller 打包后
                # 打包后，可执行文件通常在项目根目录或其子目录中
                exe_dir = os.path.dirname(sys.executable)
                # 检查是否在 python 子目录中，如果是则回到上级目录
                if os.path.basename(exe_dir) == 'python':
                    project_root = os.path.dirname(exe_dir)
                else:
                    project_root = exe_dir
            else:
                # 脚本路径: /项目根/python/matplotlibphoto.py
                # 项目根路径: /项目根
                script_dir = os.path.dirname(os.path.abspath(__file__))
                project_root = os.path.dirname(script_dir)
            
            # 构建完整路径
            full_path = os.path.join(project_root, file_path)
            full_path = os.path.abspath(full_path).replace('/', os.sep)
            
            print(f"  可执行文件/脚本目录: {os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))}")
            print(f"  项目根目录: {project_root}")
            print(f"  相对路径: {file_path}")
            print(f"  完整路径: {full_path}")
            print(f"  文件是否存在: {os.path.exists(full_path)}")
            return full_path
        
        # 处理相对路径 (如 ./src/img/xxx.png)
        elif img_src.startswith('./src/img/'):
            # 获取项目根目录
            if getattr(sys, 'frozen', False):  # PyInstaller 打包后
                exe_dir = os.path.dirname(sys.executable)
                if os.path.basename(exe_dir) == 'python':
                    project_root = os.path.dirname(exe_dir)
                else:
                    project_root = exe_dir
            else:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                project_root = os.path.dirname(script_dir)
            
            # 构建完整路径
            relative_path = img_src[2:]  # 去掉 './'
            full_path = os.path.join(project_root, relative_path)
            full_path = os.path.abspath(full_path)
            print(f"  项目根目录: {project_root}")
            print(f"  完整路径: {full_path}")
            print(f"  文件是否存在: {os.path.exists(full_path)}")
            return full_path
        
        # 处理绝对路径
        elif os.path.isabs(img_src):
            print(f"  绝对路径: {img_src}")
            print(f"  文件是否存在: {os.path.exists(img_src)}")
            return img_src
        
        # 其他情况，尝试作为相对路径处理
        else:
            if getattr(sys, 'frozen', False):
                exe_dir = os.path.dirname(sys.executable)
                if os.path.basename(exe_dir) == 'python':
                    project_root = os.path.dirname(exe_dir)
                else:
                    project_root = exe_dir
            else:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                project_root = os.path.dirname(script_dir)
            
            full_path = os.path.join(project_root, img_src)
            full_path = os.path.abspath(full_path)
            print(f"  其他情况处理 - 完整路径: {full_path}")
            print(f"  文件是否存在: {os.path.exists(full_path)}")
            return full_path
            
    except Exception as e:
        print(f"❌ 处理图片路径失败: {e}")
        return None

def extract_image_width(style_str):
    """
    从样式字符串中提取图片宽度
    """
    try:
        if 'width:' in style_str:
            # 提取宽度值，如 "width: 300px"
            width_part = [s.strip() for s in style_str.split(';') if 'width:' in s][0]
            width_value = width_part.split(':')[1].strip()
            
            if width_value.endswith('px'):
                # 将像素转换为英寸 (假设 96 DPI)
                px_value = int(width_value[:-2])
                inches = px_value / 96.0
                return Inches(min(inches, 6.0))  # 限制最大宽度为6英寸
        
        # 默认宽度
        return Inches(3.0)
    except:
        return Inches(3.0)

def generate_docx_to_pdf_base64(questions: list, soffice_path: str) -> str:
    """
    生成 Word 文档，转换为 PDF，再转换为图片，返回图片的 base64 编码。
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")

        # ✅ 生成 Word 文档（预览模式不传入文件路径，不显示元信息）
        generate_docx(questions, docx_path)

        # ✅ 使用 LibreOffice 转换成 PDF
        result = subprocess.run([
            soffice_path,
            "--headless", "--convert-to", "pdf",
            "--outdir", tmpdir,
            docx_path
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode != 0:
            print(f"❌ LibreOffice 转换失败: {result.stderr.decode()}")
            sys.exit(1)

        # ✅ 将 PDF 转换为图片并返回 base64 编码
        return convert_pdf_to_images_base64(pdf_path)

def convert_pdf_to_images_base64(pdf_path: str) -> str:
    """
    将PDF转换为图片，返回所有页面图片的base64编码（JSON格式）
    """
    try:
        # 打开PDF文档
        pdf_document = fitz.open(pdf_path)
        images_base64 = []
        
        print(f"PDF总页数: {len(pdf_document)}")
        
        for page_num in range(len(pdf_document)):
            # 获取页面
            page = pdf_document.load_page(page_num)
            
            # 设置渲染参数，提高清晰度
            zoom = 2  # 缩放倍数，提高图片质量
            mat = fitz.Matrix(zoom, zoom)
            
            # 渲染页面为图片
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为PIL Image
            img_data = pix.tobytes("png")
            
            # 转换为base64
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            images_base64.append({
                'page': page_num + 1,
                'data': img_base64
            })
            
            print(f"✅ 第 {page_num + 1} 页转换完成")
        
        pdf_document.close()
        
        # 返回JSON格式的图片数据
        import json
        return json.dumps(images_base64, ensure_ascii=False)
        
    except Exception as e:
        print(f"❌ PDF转图片失败: {e}")
        # 如果转换失败，返回空的JSON数组
        return "[]"

def get_soffice_path():
    if getattr(sys, 'frozen', False):  # PyInstaller 打包后运行
        base_dir = os.path.dirname(sys.executable)
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))

    # ➕ 回到项目根目录，然后拼接 LibreOfficePortable 路径
    root_dir = os.path.abspath(os.path.join(base_dir, "..", ".."))
    return os.path.join(root_dir, "LibreOfficePortable", "App", "libreoffice", "program", "soffice.exe")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：")
        print("  👉 生成 Word：python export_exam.py questions.json output.docx")
        print("  👉 生成答案：python export_exam.py questions.json answer output.docx")
        print("  👉 生成预览：python export_exam.py questions.json preview [base64_output.txt]")
        sys.exit(1)

    input_json = sys.argv[1]
    mode_or_output = sys.argv[2]

    # 使用加密读取功能
    try:
        questions = read_encrypted_file(input_json)
        print(f"成功读取 {len(questions)} 个题目")
    except Exception as e:
        print(f"❌ 读取输入文件失败：{e}")
        sys.exit(1)

    if mode_or_output == 'preview':
        # ✔️ 进入预览模式
        base64_output_path = sys.argv[3] if len(sys.argv) >= 4 else None
        soffice_path = get_soffice_path()
        if not os.path.exists(soffice_path):
            print(f"❌ LibreOffice 可执行文件未找到: {soffice_path}")
            sys.exit(1)

        # 修改：现在返回的是图片的base64数据（JSON格式）
        images_base64_json = generate_docx_to_pdf_base64(questions, soffice_path)

        if base64_output_path:
            with open(base64_output_path, "w", encoding="utf-8") as f:
                f.write(images_base64_json)
            print(f"✅ 图片预览 base64 已写入: {base64_output_path}")
        else:
            print(images_base64_json)
    elif mode_or_output == 'answer':
        # ✔️ 进入答案生成模式
        if len(sys.argv) < 4:
            print("❌ 答案模式需要指定输出文件")
            sys.exit(1)
        output_docx = sys.argv[3]
        generate_docx(questions, output_docx, input_json, mode='answer')
    else:
        # ✔️ 正常导出 Word（传入输入文件路径以获取元信息）
        output_docx = mode_or_output
        generate_docx(questions, output_docx, input_json)



