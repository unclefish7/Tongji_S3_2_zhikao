import json
import sys
import os
from docx import Document
from bs4 import BeautifulSoup, NavigableString
from matplotlib import rcParams
import matplotlib.pyplot as plt

rcParams['font.sans-serif'] = ['SimHei']
rcParams['axes.unicode_minus'] = False


def generate_latex_image(latex_content, image_path):
    try:
        fig, ax = plt.subplots()
        ax.axis('off')
        latex_content = latex_content.replace('\\pix', r'\pi x')
        text = ax.text(0.5, 0.5, f'${latex_content}$', fontsize=20, ha='center', va='center')
        fig.canvas.draw()
        bbox = text.get_window_extent(renderer=fig.canvas.get_renderer())
        width_inch = bbox.width / fig.dpi
        height_inch = bbox.height / fig.dpi
        fig.set_size_inches(width_inch, height_inch)
        image_file_path = os.path.join(image_path, 'latex_image.png')
        plt.savefig(image_file_path, dpi=300, bbox_inches='tight', pad_inches=0.05)
        plt.close(fig)
        return image_file_path
    except Exception as e:
        print(f"Error generating LaTeX image: {e}")
        return None

def insert_local_image(doc, img_path):
    with open(img_path, 'rb') as f:
        doc.add_paragraph().add_picture(f)

def insert_table_from_html(doc, table_html):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(table_html, "html.parser")
    table_tag = soup.find('table')
    if not table_tag:
        return
    rows = table_tag.find_all('tr')
    if not rows:
        return
    num_cols = len(rows[0].find_all(['th', 'td']))
    table = doc.add_table(rows=0, cols=num_cols)
    for row in rows:
        cells = row.find_all(['th', 'td'])
        row_cells = table.add_row().cells
        for i in range(min(len(cells), num_cols)):
            row_cells[i].text = cells[i].get_text(strip=True)

def process_rich_text_for_answer(doc, content, is_subjective=False):
    soup = BeautifulSoup(content, 'html.parser')
    latex_image_path = os.path.join(os.getcwd(), 'images')
    os.makedirs(latex_image_path, exist_ok=True)

    for elem in soup.children:
        if isinstance(elem, NavigableString):
            continue
        if elem.name == 'p':
            if elem.find('span', attrs={"data-w-e-type": "formula"}):
                span = elem.find('span', attrs={"data-w-e-type": "formula"})
                latex_code = span.get('data-value', '')
                image_path = generate_latex_image(latex_code, latex_image_path)
                if image_path:
                    doc.add_picture(image_path)
            elif elem.find('img'):
                img_url = elem.find('img').get('src', '')
                if img_url.startswith("file://"):
                    img_path = img_url.replace("file://", ".")
                    insert_local_image(doc, img_path)
            else:
                text = elem.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
        elif elem.name == 'table':
            insert_table_from_html(doc, str(elem))

    if is_subjective:
        for _ in range(3):
            doc.add_paragraph('\n')  # 插入空行作为作答区域

def generate_answer_sheet(questions, output_path):
    doc = Document()
    doc.add_paragraph('研究生招生考试 - 答题卡')
    doc.add_paragraph('准考证号：_______________ 姓名：_______________')
    doc.add_paragraph('考试类别：[  ] 平时测验    [  ] 期中测试    [  ] 期末测试')
    doc.add_paragraph('--------------------------------------------')

    sections = {
        '选择题': [],
        '判断题': [],
        '填空题': [],
        '主观题': []
    }

    for q in questions:
        qtype = q.get('type', '').strip()
        if qtype in sections:
            sections[qtype].append(q)

    for title, items in sections.items():
        if items:
            doc.add_paragraph(f'\n{title}')
            for idx, q in enumerate(items, 1):
                doc.add_paragraph(f'{idx}.')
                if title in ['选择题', '判断题','填空题']:
                    doc.add_paragraph('答：______________________________')
                elif title in ['主观题']:
                    for _ in range(6):  # 留出更多空白区域
                        doc.add_paragraph()

    doc.save(output_path)

if __name__ == "__main__":
    json_file = sys.argv[1]
    docx_file = sys.argv[2]

    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_answer_sheet(data, docx_file)
    print(f"Answer sheet saved to {docx_file}")
